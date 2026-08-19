"""
Generates two comprehensive demo files for the Import Word/Excel feature:
  - sample-survey.xlsx — covers every supported question type + every config column
  - sample-survey.docx — covers every supported question type + every directive

Run: python3 scripts/build_samples.py
Output: public/samples/sample-survey.{xlsx,docx}
"""
from pathlib import Path
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill, Border, Side
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


HERE = Path(__file__).resolve().parent.parent
OUT = HERE / "public" / "samples"
OUT.mkdir(parents=True, exist_ok=True)

# ---------------- XLSX ----------------
def build_xlsx():
    wb = Workbook()
    ws = wb.active
    ws.title = "Questions"

    headers = [
        "Type", "Question", "Options", "Required",
        "Rows", "Cols", "Min", "Max", "Scale", "Placeholder",
        "Skip If", "Skip Combinator", "Randomize", "Piping Fallback",
    ]
    # Q-references in Skip-If are 1-indexed within THIS import batch.
    # Q1 = first row, Q2 = second row, etc.
    rows = [
        # ---- Q1 — Standard single (drives skip + piping) ----
        ["single",
         "How often did you shop with Northwind Retail in the past 30 days?",
         "Daily; 2-3 times per week; Weekly; Monthly; Rarely / Never",
         "Y", "", "", "", "", "", "",
         "", "", "", ""],

        # ---- Q2 — Multi-select with RANDOMIZED options + skip if Q1 = Never ----
        ["multi",
         "Which Northwind services do you use? (Select all that apply)",
         "In-store shopping; Online orders; Click & collect; Home delivery; Gift cards; Loyalty programme; Returns & exchanges; Personal shopper; Subscription box",
         "N", "", "", "", "", "", "",
         "Q1!=Rarely / Never", "AND", "random", ""],

        # ---- Q3 — Dropdown (no rules, just standard) ----
        ["dropdown",
         "Which division do you live in?",
         "North; South; East; West; Central; Metro; Coastal; Highlands",
         "Y", "", "", "", "", "", "",
         "", "", "", ""],

        # ---- Q4 — Matrix with RANDOMIZED rows + skip if Q1 = Never ----
        ["matrix",
         "How would you rate each of the following aspects of Northwind Retail?",
         "", "N",
         "Checkout speed; Product quality; Customer service; Loyalty rewards; Store locations; Ease of returns",
         "Poor; Fair; Good; Very good; Excellent",
         "", "", "", "",
         "Q1!=Rarely / Never", "AND", "random", ""],

        # ---- Q5 — Rating (used to gate Q8 + Q14 below) ----
        ["rating",
         "Overall, how satisfied are you with Northwind Retail?",
         "", "Y", "", "",
         "", "", "5", "",
         "Q1!=Rarely / Never", "AND", "", ""],

        # ---- Q6 — Slider, skipped for non-users ----
        ["slider",
         "Roughly how much do you spend with Northwind Retail per month?",
         "", "N", "", "",
         "0", "50000", "", "",
         "Q1!=Rarely / Never", "AND", "", ""],

        # ---- Q7 — Rank with RANDOMIZED items ----
        ["rank",
         "Rank these features by importance. (drag to reorder)",
         "Faster checkout; Lower prices; Better product quality; More loyalty rewards; More store locations",
         "N", "", "", "", "", "", "",
         "", "", "random", ""],

        # ---- Q8 — Open text WITH PIPING + skip if rated low OR NPS is detractor ----
        ["text",
         "You rated us {{Q5}}/5. What's the one thing we could improve?",
         "", "N", "", "",
         "", "500", "",
         "Tell us in your own words…",
         "Q5<=3; Q15<=6", "OR", "", "thanks for your honest feedback"],

        # ---- Q9 — Video stimulus ----
        ["video",
         "Watch this short clip and share your reaction.",
         "https://commondatastorage.googleapis.com/gtv-videos-bucket/sample/ForBiggerBlazes.mp4",
         "N", "", "", "", "", "", "",
         "", "", "", ""],

        # ---- Q10 — Hotspot ----
        ["hotspot",
         "Click on the area of the home screen that you find most useful.",
         "https://placehold.co/600x400/png?text=Northwind+store+layout",
         "N", "", "", "", "", "", "",
         "", "", "", ""],

        # ---- Q11 — Image options with PIN-LAST ----
        ["image-opt",
         "Which colour palette feels most like Northwind to you?",
         "🟪; 🟦; 🟩; 🟧; ⬛",
         "N", "", "", "", "", "", "",
         "", "", "pin-last", ""],

        # ---- Q12 — IAT ----
        ["iat",
         "Quick! React as fast as you can — is Northwind positive or negative?",
         "Positive; Negative",
         "N", "", "", "", "", "", "",
         "", "", "", ""],

        # ---- Q13 — Fill in the blanks ----
        ["fill",
         "Complete the sentence: Shopping with Northwind makes my week ___ and the one thing I miss is ___.",
         "", "N", "", "", "", "", "", "",
         "", "", "", ""],

        # ---- Q14 — Audio (only if low rating OR low NPS) ----
        ["audio",
         "Record a short voice message for our product team.",
         "", "N", "", "", "", "30", "", "",
         "Q5<=3; Q15<=6", "OR", "", ""],

        # ---- Q15 — NPS WITH PIPING (references Q1) ----
        ["nps",
         "You said you're {{Q1}} — how likely are you to recommend Northwind Retail to a friend?",
         "", "Y", "", "",
         "", "", "11", "",
         "", "", "", "a customer"],

        # ---- Q16 — Constant Sum ----
        ["constant-sum",
         "Allocate 100 points across these priorities — give more to what matters most.",
         "Checkout speed; Loyalty rewards; Product quality; Customer service; Store locations",
         "N", "", "", "", "", "", "",
         "", "", "", ""],
    ]

    # Header styling
    bold = Font(bold=True, color="FFFFFFFF", name="Calibri", size=11)
    fill_head = PatternFill("solid", fgColor="FF1F2937")
    border = Border(
        left=Side(style="thin", color="FFE5E7EB"),
        right=Side(style="thin", color="FFE5E7EB"),
        top=Side(style="thin", color="FFE5E7EB"),
        bottom=Side(style="thin", color="FFE5E7EB"),
    )

    for col_idx, h in enumerate(headers, start=1):
        c = ws.cell(row=1, column=col_idx, value=h)
        c.font = bold
        c.fill = fill_head
        c.alignment = Alignment(horizontal="left", vertical="center")
        c.border = border

    # Data rows
    for r_idx, row in enumerate(rows, start=2):
        for c_idx, val in enumerate(row, start=1):
            c = ws.cell(row=r_idx, column=c_idx, value=val)
            c.alignment = Alignment(vertical="top", wrap_text=True)
            c.border = border
            if c_idx == 1:
                c.font = Font(name="Consolas", size=10, bold=True, color="FFE2136B")

    # Column widths
    widths = [12, 60, 60, 10, 35, 35, 8, 10, 8, 30,
              28, 18, 14, 30]
    for i, w in enumerate(widths, start=1):
        ws.column_dimensions[get_column_letter(i)].width = w

    # Row heights — give breathing room
    ws.row_dimensions[1].height = 22
    for r in range(2, len(rows) + 2):
        ws.row_dimensions[r].height = 36
    ws.freeze_panes = "A2"

    # ----- Reference / legend sheet -----
    ref = wb.create_sheet("README")
    ref["A1"] = "Consumer Insights Platform — Excel Import Reference"
    ref["A1"].font = Font(bold=True, size=14)

    ref["A3"] = "Sheet rules"
    ref["A3"].font = Font(bold=True, size=12)
    rules = [
        "• First sheet must contain the questions. Sheet name does not matter.",
        "• First row is the header row. Column order does not matter — columns are matched by name.",
        "• Required headers: Type, Question. All others are optional.",
        "• Blank rows are skipped. A row with no Type or Question is treated as blank.",
        "• Multi-value cells use ';' (preferred), '|' or new-line as separators.",
        "• 'Required' accepts: Y, Yes, true, 1, required (case-insensitive). Anything else = optional.",
        "• 'Randomize' values: random  |  pin-last  |  (blank for no shuffle).",
        "• 'Skip If' clause syntax:  Q1=Daily   Q1!=Never   Q5<=3   Q5>=8   Q1~good   Q1=  (empty).",
        "    Multiple clauses are separated by ';'. Q references are 1-indexed within this import batch.",
        "• 'Skip Combinator' is AND or OR (default AND, applies when multiple clauses are listed).",
        "• 'Piping Fallback' is the string used at runtime when a piped answer is missing.",
        "• Piping syntax in the Question column itself: {{Q1}} or {{Q1.answer}}.",
    ]
    for i, line in enumerate(rules):
        ref.cell(row=4 + i, column=1, value=line)

    ref["A12"] = "Question types"
    ref["A12"].font = Font(bold=True, size=12)
    type_doc = [
        ("single",        "Single-select radio. Uses Options."),
        ("multi",         "Multi-select checkbox. Uses Options."),
        ("dropdown",      "Compact select list. Uses Options."),
        ("matrix",        "Grid: Rows × Cols. Both required."),
        ("rating",        "Star/scale rating. Scale = 3 / 5 / 7 / 10."),
        ("slider",        "Continuous slider. Min, Max, optional Scale = step."),
        ("rank",          "Drag-and-drop ranking. Items via Options."),
        ("text",          "Open text. Optional Placeholder, Min/Max char counts via Min/Max."),
        ("video",         "Embedded video. URL goes into Options."),
        ("hotspot",       "Image with click-region capture. URL goes into Options."),
        ("image-opt",     "Visual options (emoji or image URLs). Uses Options."),
        ("iat",           "Implicit Association Test — captures reaction time. Two options."),
        ("fill",          "Fill in the blanks. Use ___ as a blank in the question text."),
        ("audio",         "Voice response. Max length goes into Max (seconds)."),
        ("nps",           "Net Promoter Score 0-10. Scale defaults to 11."),
        ("constant-sum",  "Allocate N points across items (Options). Total defaults to 100."),
    ]
    ref["A13"] = "type"
    ref["B13"] = "description"
    ref["A13"].font = Font(bold=True)
    ref["B13"].font = Font(bold=True)
    for i, (t, d) in enumerate(type_doc):
        ref.cell(row=14 + i, column=1, value=t).font = Font(name="Consolas", size=10)
        ref.cell(row=14 + i, column=2, value=d)

    ref.column_dimensions["A"].width = 18
    ref.column_dimensions["B"].width = 80

    out = OUT / "sample-survey.xlsx"
    wb.save(out)
    return out


# ---------------- DOCX ----------------
def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt({1: 18, 2: 14, 3: 12}.get(level, 11))
    if level == 1:
        run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
    return p


def add_mono(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
    return p


def add_question(doc, tag, text, options=None, rows=None, cols=None,
                 required=False, placeholder=None, scale=None,
                 mn=None, mx=None, total=None, url=None,
                 skip_if=None, skip_combinator=None,
                 randomize=None, piping_fallback=None):
    """Writes one question paragraph block in the [TAG] format the parser expects."""
    p = doc.add_paragraph()
    run = p.add_run(f"[{tag.upper()}] {text}")
    run.bold = True
    run.font.color.rgb = RGBColor(0xE2, 0x13, 0x6B)

    if required:
        doc.add_paragraph("REQUIRED: yes")
    if placeholder:
        doc.add_paragraph(f"PLACEHOLDER: {placeholder}")
    if scale is not None:
        doc.add_paragraph(f"SCALE: {scale}")
    if mn is not None:
        doc.add_paragraph(f"MIN: {mn}")
    if mx is not None:
        doc.add_paragraph(f"MAX: {mx}")
    if total is not None:
        doc.add_paragraph(f"TOTAL: {total}")
    if url:
        doc.add_paragraph(f"URL: {url}")
    if rows:
        doc.add_paragraph(f"ROWS: {'; '.join(rows)}")
    if cols:
        doc.add_paragraph(f"COLS: {'; '.join(cols)}")
    if randomize:
        doc.add_paragraph(f"RANDOMIZE: {randomize}")
    if piping_fallback:
        doc.add_paragraph(f"PIPING FALLBACK: {piping_fallback}")
    if skip_combinator:
        doc.add_paragraph(f"SKIP COMBINATOR: {skip_combinator}")
    if skip_if:
        doc.add_paragraph(f"SKIP IF: {skip_if}")
    if options:
        for opt in options:
            doc.add_paragraph(f"- {opt}")
    doc.add_paragraph("")  # blank separator


def build_docx():
    doc = Document()
    # default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ---- Cover ----
    add_heading(doc, "Consumer Insights Platform", level=1)
    add_heading(doc, "Word Import Reference Document", level=2)
    p = doc.add_paragraph(
        "Save this file as .docx and upload it via Builder → Import Word/Excel. "
        "Each [TAG] paragraph defines one question; lines starting with '- ' are options; "
        "ROWS:/COLS:/REQUIRED:/PLACEHOLDER:/SCALE:/MIN:/MAX:/TOTAL:/URL: directives configure the question."
    )
    p.runs[0].font.size = Pt(10)
    doc.add_paragraph("")

    add_heading(doc, "Format rules", level=2)
    rules = [
        "• Each question begins with a [TYPE] tag followed by the question text on the same line.",
        "• Aliases work: [SINGLE]/[RADIO]/[MCQ], [MULTI]/[CHECKBOX], [TEXT]/[OPEN-ENDED], [RATING]/[LIKERT], [RANK]/[RANKING].",
        "• Options: lines starting with '-', '*' or '•'.",
        "• Matrix questions use 'ROWS:' and 'COLS:' on their own lines, '; '-separated.",
        "• 'REQUIRED: yes' makes the question required (yes / y / true / 1).",
        "• 'PLACEHOLDER:' sets the empty-state hint for text inputs.",
        "• 'SCALE:' sets the rating/NPS scale (NPS defaults to 11, rating to 5).",
        "• 'MIN:'/'MAX:' configure sliders, char limits, audio max-seconds.",
        "• 'TOTAL:' sets the constant-sum total (defaults to 100).",
        "• 'URL:' sets the media URL for video/hotspot/image-opt questions.",
        "• 'RANDOMIZE:' sets question option/item shuffling.  Values: random | pin-last | none.",
        "• 'SKIP IF:' shows the question only if all (AND) or any (OR) clauses match.",
        "    Clause syntax:  Q1=Daily   Q1!=Never   Q5<=3   Q5>=8   Q1~good   Q1=  (empty)",
        "    Multiple clauses are separated by ';'. Default combinator is AND.",
        "• 'SKIP COMBINATOR: OR' makes multi-clause SKIP IF rules use OR instead of AND.",
        "• Piping in question text: write {{Q1}} or {{Q1.answer}} — substituted at runtime.",
        "• 'PIPING FALLBACK:' sets the string to substitute if the source answer is empty.",
        "• Blank lines between questions are optional but recommended for readability.",
    ]
    for r in rules:
        doc.add_paragraph(r)
    doc.add_paragraph("")

    add_heading(doc, "Worked examples — every question type", level=2)
    doc.add_paragraph(
        "Copy the block below into your own document, edit the text, and upload."
    )

    # ---- 1 single ----
    add_question(doc, "SINGLE",
                 "How often did you shop with Northwind Retail in the past 30 days?",
                 options=["Daily", "2-3 times per week", "Weekly", "Monthly", "Rarely / Never"],
                 required=True)

    # ---- 2 multi (RANDOMIZED + skipped if Q1 = Never) ----
    add_question(doc, "MULTI",
                 "Which Northwind services do you use? (Select all that apply)",
                 options=[
                     "In-store shopping", "Online orders", "Click & collect", "Home delivery",
                     "Gift cards", "Loyalty programme", "Returns & exchanges", "Personal shopper",
                     "Subscription box",
                 ],
                 randomize="random",
                 skip_if="Q1!=Rarely / Never")

    # ---- 3 dropdown ----
    add_question(doc, "DROPDOWN",
                 "Which division do you live in?",
                 options=[
                     "North", "South", "East", "West",
                     "Central", "Metro", "Coastal", "Highlands",
                 ],
                 required=True)

    # ---- 4 matrix (RANDOMIZED rows + skipped for non-users) ----
    add_question(doc, "MATRIX",
                 "How would you rate each of the following aspects of Northwind Retail?",
                 rows=["Checkout speed", "Product quality", "Customer service",
                       "Loyalty rewards", "Store locations", "Ease of returns"],
                 cols=["Poor", "Fair", "Good", "Very good", "Excellent"],
                 randomize="random",
                 skip_if="Q1!=Rarely / Never")

    # ---- 5 rating (skipped for non-users; used to gate Q8 + Q14) ----
    add_question(doc, "RATING",
                 "Overall, how satisfied are you with Northwind Retail?",
                 required=True, scale=5,
                 skip_if="Q1!=Rarely / Never")

    # ---- 6 slider (skipped for non-users) ----
    add_question(doc, "SLIDER",
                 "Roughly how much do you spend with Northwind Retail per month?",
                 mn=0, mx=50000,
                 skip_if="Q1!=Rarely / Never")

    # ---- 7 rank (RANDOMIZED items) ----
    add_question(doc, "RANK",
                 "Rank these features by importance. (drag to reorder)",
                 options=[
                     "Faster checkout", "Lower prices", "Better product quality",
                     "More loyalty rewards", "More store locations",
                 ],
                 randomize="random")

    # ---- 8 text — PIPING + skip if low rating OR low NPS ----
    add_question(doc, "TEXT",
                 "You rated us {{Q5}}/5. What's the one thing we could improve?",
                 placeholder="Tell us in your own words…",
                 mx=500,
                 piping_fallback="thanks for your honest feedback",
                 skip_combinator="OR",
                 skip_if="Q5<=3; Q15<=6")

    # ---- 9 video ----
    add_question(doc, "VIDEO",
                 "Watch this short clip and share your reaction.",
                 url="https://commondatastorage.googleapis.com/gtv-videos-bucket/sample/ForBiggerBlazes.mp4")

    # ---- 10 hotspot ----
    add_question(doc, "HOTSPOT",
                 "Click on the area of the home screen that you find most useful.",
                 url="https://placehold.co/600x400/png?text=Northwind+store+layout")

    # ---- 11 image-opt (PIN-LAST randomization) ----
    add_question(doc, "IMAGE-OPT",
                 "Which colour palette feels most like Northwind to you?",
                 options=["🟪", "🟦", "🟩", "🟧", "⬛"],
                 randomize="pin-last")

    # ---- 12 iat ----
    add_question(doc, "IAT",
                 "Quick! React as fast as you can — is Northwind positive or negative?",
                 options=["Positive", "Negative"])

    # ---- 13 fill ----
    add_question(doc, "FILL",
                 "Complete the sentence: Shopping with Northwind makes my week ___ and the one thing I miss is ___.")

    # ---- 14 audio (skipped unless rating ≤ 3 OR NPS ≤ 6) ----
    add_question(doc, "AUDIO",
                 "Record a short voice message for our product team.",
                 mx=30,
                 skip_combinator="OR",
                 skip_if="Q5<=3; Q15<=6")

    # ---- 15 nps — PIPING references Q1 ----
    add_question(doc, "NPS",
                 "You said you're {{Q1}} — how likely are you to recommend Northwind Retail to a friend?",
                 required=True, scale=11,
                 piping_fallback="a customer")

    # ---- 16 constant-sum ----
    add_question(doc, "CONSTANT-SUM",
                 "Allocate 100 points across these priorities — give more to what matters most.",
                 options=["Checkout speed", "Loyalty rewards", "Product quality", "Customer service", "Store locations"],
                 total=100)

    add_heading(doc, "Tips", level=2)
    tips = [
        "• You can mix and match — only TAG and the question text are required, everything else is optional.",
        "• Free-floating text before the first [TAG] is ignored, so feel free to add cover notes.",
        "• Lines that don't match a directive get appended to the previous question's text.",
        "• Re-running the importer is safe — every import appends to the end of your survey, never overwrites.",
    ]
    for t in tips:
        doc.add_paragraph(t)

    out = OUT / "sample-survey.docx"
    doc.save(out)
    return out


if __name__ == "__main__":
    xlsx_path = build_xlsx()
    docx_path = build_docx()
    print(f"Wrote {xlsx_path}")
    print(f"Wrote {docx_path}")
